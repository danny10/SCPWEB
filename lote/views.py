# encoding: utf-8
from django.shortcuts import render,redirect,render_to_response
from django.core.urlresolvers import reverse
from rest_framework import viewsets, serializers
from django.db.models import Q
from django.template import RequestContext
from rest_framework.renderers import JSONRenderer
from rest_framework.response import Response
from rest_framework.request import Request
from rest_framework import status
from rest_framework.pagination import PageNumberPagination
import json
from .models import Lote,LoteCircuito,LoteSector,LotePoligono,Mes
from inspeccion.models import AInspeccion,BFallaInspeccion,FotoFallaInspeccion
from provincia.models import Provincia
from contratista.models import Contratista
from contratista.views import ContratistaSerializer
from provincia.models import Provincia
from provincia.views import ProvinciaSerializer
from sucursal.models import Sucursal
from poligono.models import Poligono
from circuito.models import Circuito
from sector.models import Sector
from sucursal.views import SucursalSerializer
from django.http import HttpResponse,JsonResponse
from django.db import transaction,connection
from django.db.models.deletion import ProtectedError
from django.contrib.auth.decorators import login_required
from rest_framework.decorators import api_view, throttle_classes
from SCPWEB.functions import functions
from docx import Document

from docx.shared import Inches , Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH , WD_COLOR_INDEX

from docx.section import Section, Sections

from docx.enum.section import WD_ORIENT


import os
import mimetypes
from django.http import StreamingHttpResponse
from django.core.servers.basehttp import FileWrapper


#Api para los lotes
class LoteSerializer(serializers.HyperlinkedModelSerializer):

	contratista_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Contratista.objects.all())
	contratista=ContratistaSerializer(read_only=True)

	provincia_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Provincia.objects.all())
	provincia=ProvinciaSerializer(read_only=True)

	sucursal_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Sucursal.objects.all())
	sucursal=SucursalSerializer(read_only=True)

	lotecircuito = serializers.SerializerMethodField()
	lotesector = serializers.SerializerMethodField()

	class Meta:
		model = Lote
		fields=('id','contratista','contratista_id','provincia','provincia_id','sucursal','sucursal_id','nombre','activo','lotecircuito','lotesector')


	def get_lotecircuito(self, obj):

		listado=LoteCircuito.objects.filter(lote__id=obj.id)
		listcircuito=''

		for item in list(listado):

			listcircuito='{}{},'.format(listcircuito,item.circuito.nombre)

		return listcircuito


	def get_lotesector(self, obj):

		listado2=LoteSector.objects.filter(lote__id=obj.id)
		listsector=''

		for item in list(listado2):

			listsector='{}{},'.format(listsector,item.sector.nombre)

		return listsector


class LoteViewSet(viewsets.ModelViewSet):
	""" Retorna una lista de lotes. """
	model=Lote
	queryset = model.objects.all()
	serializer_class = LoteSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)


	def list(self, request, *args, **kwargs):
		try:

			queryset = super(LoteViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			contratista= self.request.query_params.get('id_contratista',None)
			provincia= self.request.query_params.get('id_provincia',None)
			sucursal= self.request.query_params.get('id_sucursal',None)
			circuito= self.request.query_params.get('id_circuito',None)
			sector= self.request.query_params.get('id_sector',None)
			activo= self.request.query_params.get('activo',None)
			sin_paginacion=self.request.query_params.get('sin_paginacion',None)
			qset=''

			qset = Q(activo=activo)

			if (dato or contratista or provincia or sucursal or circuito):
			

				if dato:
					qset = qset &(
						Q(nombre__icontains=dato)
					)

				if contratista:
					qset = qset &(
						Q(contratista__id=contratista)
					)


				if provincia:
					qset = qset &(
						Q(provincia__id=provincia)
					)


				if sucursal:
					qset = qset &(
						Q(sucursal__id=sucursal)
					)

				if circuito:
					qset = qset &(
						Q(id__in=LoteCircuito.objects.filter(circuito__id=circuito).values('lote__id'))
					)


				if sector:
					qset = qset &(
						Q(id__in=LoteSector.objects.filter(sector__id=sector).values('lote__id'))
					)

			#print qset
			if qset != '':
				queryset = self.model.objects.filter(qset)

			if sin_paginacion is None:	
	
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					#print serializer
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})
				
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok',
						'data':serializer.data})
			
			else:
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok',
					'data':serializer.data})
											
		except Exception,e:
			print e
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()

			try:
				serializer = LoteSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():
					serializer.save(contratista_id=request.DATA['contratista_id'],provincia_id=request.DATA['provincia_id'],sucursal_id=request.DATA['sucursal_id'])

					transaction.savepoint_commit(sid)

					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					print serializer.errors
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

			except Exception,e:
				#print e
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)	


	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()

			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = LoteSerializer(instance,data=request.DATA,context={'request': request},partial=partial)

				if serializer.is_valid():
					serializer.save(contratista_id=request.DATA['contratista_id'],provincia_id=request.DATA['provincia_id'],sucursal_id=request.DATA['sucursal_id'])

					transaction.savepoint_commit(sid)

					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
				 	return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except:

				transaction.savepoint_rollback(sid)
			 	return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

	def destroy(self,request,*args,**kwargs):
		try:
			instance = self.get_object()
			self.perform_destroy(instance)
			return Response({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''},status=status.HTTP_204_NO_CONTENT)
		except:
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''},status=status.HTTP_400_BAD_REQUEST)	


#Fin api  para lotes.



#Api para los poligonos
class PoligonoSerializer(serializers.HyperlinkedModelSerializer):

	class Meta:
		model = Poligono
		fields=('id','nombre','activo')


class PoligonoViewSet(viewsets.ModelViewSet):
	""" Retorna una lista de poligonos. """
	model=Poligono
	queryset = model.objects.all()
	serializer_class = PoligonoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)


	def list(self, request, *args, **kwargs):
		try:

			queryset = super(PoligonoViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			sin_paginacion = self.request.query_params.get('sin_paginacion', None)
			qset=''

			if dato:

				qset = Q(nombre__icontains=dato)

			#print qset
			if qset != '':
				queryset = self.model.objects.filter(qset)

			if sin_paginacion is None:	
	
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})
				
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok','data':serializer.data})
			
			else:
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok','data':serializer.data})
											
		except Exception,e:
			#print e
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()

			try:
				serializer = PoligonoSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():

					serializer.save()

					poligonoLote=LotePoligono(lote_id=request.DATA['lote_id'],poligono_id=serializer.data['id'])
					poligonoLote.save()

					transaction.savepoint_commit(sid)

					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					#print serializer.errors
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

			except Exception,e:
				print e
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)	


	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()

			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = PoligonoSerializer(instance,data=request.DATA,context={'request': request},partial=partial)

				if serializer.is_valid():
					serializer.save()

					transaction.savepoint_commit(sid)

					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
				 	return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except:

				transaction.savepoint_rollback(sid)
			 	return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

	def destroy(self,request,*args,**kwargs):
		try:
			instance = self.get_object()
			self.perform_destroy(instance)
			return Response({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''},status=status.HTTP_204_NO_CONTENT)
		except:
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''},status=status.HTTP_400_BAD_REQUEST)	

#Fin de api para poligonos



#Api para los lote poligono
class LotePoligonoSerializer(serializers.HyperlinkedModelSerializer):

	lote_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Lote.objects.all())
	lote=LoteSerializer(read_only=True)

	poligono_id = serializers.PrimaryKeyRelatedField(write_only=True,queryset=Poligono.objects.all())
	poligono=PoligonoSerializer(read_only=True)


	class Meta:
		model = LotePoligono
		fields=('id','lote','lote_id','poligono','poligono_id')


class LotePoligonoViewSet(viewsets.ModelViewSet):
	""" Retorna una lista de lotes. """
	model=LotePoligono
	queryset = model.objects.all()
	serializer_class = LotePoligonoSerializer

	def retrieve(self,request,*args, **kwargs):
		try:
			instance = self.get_object()
			serializer = self.get_serializer(instance)
			return Response({'message':'','status':'success','data':serializer.data})
		except:
			return Response({'message':'No se encontraron datos','success':'fail','data':''},status=status.HTTP_404_NOT_FOUND)


	def list(self, request, *args, **kwargs):
		try:

			queryset = super(LotePoligonoViewSet, self).get_queryset()
			dato = self.request.query_params.get('dato', None)
			lote= self.request.query_params.get('id_lote',None)
			poligono= self.request.query_params.get('id_poligono',None)
			activo= self.request.query_params.get('activo',None)
			sin_paginacion=self.request.query_params.get('sin_paginacion',None)
			qset=''

			if (dato or lote or poligono or activo):
			
				qset = Q(poligono__activo=activo)

				if lote and int(lote)>0:
					qset = qset &(
						Q(lote__id=lote)
					)


				if dato:
					qset = qset &(
						Q(poligono__nombre__icontains=dato)
					)

			#print qset
			if qset != '':
				queryset = self.model.objects.filter(qset)

			if sin_paginacion is None:	
	
				page = self.paginate_queryset(queryset)
				if page is not None:
					serializer = self.get_serializer(page,many=True)	
					#print serializer
					return self.get_paginated_response({'message':'','success':'ok',
					'data':serializer.data})
				
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok',
						'data':serializer.data})
			
			else:
				serializer = self.get_serializer(queryset,many=True)
				return Response({'message':'','success':'ok',
					'data':serializer.data})
											
		except Exception,e:
			print e
			return Response({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)


	@transaction.atomic
	def create(self, request, *args, **kwargs):
		if request.method == 'POST':
			sid = transaction.savepoint()

			try:
				serializer = LotePoligonoSerializer(data=request.DATA,context={'request': request})

				if serializer.is_valid():
					serializer.save(lote_id=request.DATA['lote_id'],poligono_id=request.DATA['poligono_id'])

					transaction.savepoint_commit(sid)

					return Response({'message':'El registro ha sido guardado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
					print serializer.errors
					return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

			except Exception,e:
				#print e
				transaction.savepoint_rollback(sid)
				return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)	


	@transaction.atomic
	def update(self,request,*args,**kwargs):
		if request.method == 'PUT':
			sid = transaction.savepoint()

			try:
				partial = kwargs.pop('partial', False)
				instance = self.get_object()
				serializer = LotePoligonoSerializer(instance,data=request.DATA,context={'request': request},partial=partial)

				if serializer.is_valid():
					serializer.save(lote_id=request.DATA['lote_id'],poligono_id=request.DATA['poligono_id'])

					transaction.savepoint_commit(sid)

					return Response({'message':'El registro ha sido actualizado exitosamente','success':'ok',
						'data':serializer.data},status=status.HTTP_201_CREATED)
				else:
				 	return Response({'message':'datos requeridos no fueron recibidos','success':'fail',
			 		'data':''},status=status.HTTP_400_BAD_REQUEST)
			except:

				transaction.savepoint_rollback(sid)
			 	return Response({'message':'Se presentaron errores al procesar los datos','success':'error',
					'data':''},status=status.HTTP_400_BAD_REQUEST)

	def destroy(self,request,*args,**kwargs):
		try:
			instance = self.get_object()
			self.perform_destroy(instance)
			return Response({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''},status=status.HTTP_204_NO_CONTENT)
		except:
			return Response({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''},status=status.HTTP_400_BAD_REQUEST)	


#Fin api  para lotes poligono.




# guarda lotes
@api_view(['POST'])
def guardar_lote(request):

	if request.method == 'POST':

		try:

			listaSector=request.DATA['lista_sector']
			listaCircuito= request.DATA['lista_circuito']
			listaPoligono= request.DATA['lote_poligono']

			modelo=Lote(nombre=request.DATA['nombre_lote'],
						contratista_id=request.DATA['contratista_id'],
						provincia_id=request.DATA['municipio_id'],
						sucursal_id=request.DATA['sucursal_id'], 
						activo=request.DATA['activo'])
			modelo.save()


			for itemcircuito in listaCircuito:

				if itemcircuito['eliminado']==False:

					modelo2=LoteCircuito(circuito_id=itemcircuito['id'],lote_id=modelo.id)
					modelo2.save()


			for itemsector in listaSector:

				if itemsector['eliminado']==False:

					modelo3=LoteSector(sector_id=itemsector['id'],lote_id=modelo.id)
					modelo3.save()


			for itempoligono in listaPoligono:

				if itempoligono['eliminado']==False:

					modelo3=Poligono(nombre=itempoligono['nombre'],activo=True)
					modelo3.save()


					modelo4=LotePoligono(poligono_id=modelo3.id,lote_id=modelo.id)
					modelo4.save()
			
			return JsonResponse({'message':'Se registro exitosamente','success':'ok','data':''})

		
		except Exception,e:
			functions.toLog(e,'lote')
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)




#Actualiza a inactivo el lote
@transaction.atomic
def LoteInactivo(request):

	sid = transaction.savepoint()

	try:
		lista=request.POST['_content']
		respuesta= json.loads(lista)

		for item in respuesta['lista']:
			object_lote=Lote.objects.get(pk=int(item['id']))
			object_lote.activo=False

			object_lote.save()

			transaction.savepoint_commit(sid)

		return JsonResponse({'message':'El registro se ha guardado correctamente','success':'ok',
				'data':''})
		
	except Exception,e:
		print e
		transaction.savepoint_rollback(sid)
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})




#consultar circuito, pagina editar
@api_view(['GET'])
def consulta_listado_circuito(request):

	if request.method == 'GET':

		try:
			qset=''
			dato= request.GET['dato']
			id_lote= request.GET['id_lote']

			qset = Q(activo=True)

			if dato:
					qset =  qset &(Q(nombre__icontains=dato))


			circuito = Circuito.objects.filter(qset)

			lista_circuito=[]

			for item in list(circuito):

			 	circuito_lote = LoteCircuito.objects.filter(circuito__id=item.id,lote__id=id_lote).first()

			 	if circuito_lote is None:

					lista_circuito.append(

						{		
							'id':item.id,
							'nombre':item.nombre,
						}

					)

			return JsonResponse({'message':'','success':'ok','data':lista_circuito})

		except Exception,e:
			print e
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

	    #return response



#consultar circuito asociados al lote, pagina editar
@api_view(['GET'])
def consulta_listado_circuito_asociados(request):

	if request.method == 'GET':

		try:

			id_lote= request.GET['id_lote']

			circuito__asociado_lote = LoteCircuito.objects.filter(lote__id=id_lote)

			lista_circuito_asociado=[]

			for item in list(circuito__asociado_lote):

				lista_circuito_asociado.append(

					{		
						'id':item.id,
						'nombre':item.circuito.nombre,
					}

				)

			return JsonResponse({'message':'','success':'ok','data':lista_circuito_asociado})

		except Exception,e:
			print e
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

    #return response



#guardar la asociacion de los circuito al lote pagina editar
@api_view(['POST'])
@transaction.atomic
def AsociarCircuitoLoteEditar(request):

	sid = transaction.savepoint()
	if request.method == 'POST':
		try:
			
			lista=request.DATA['lista']
			lote_id= request.DATA['lote_id']
  
			for item in lista:

				circuitoLote=LoteCircuito(circuito_id=item['id'],lote_id=lote_id)
				circuitoLote.save()

				transaction.savepoint_commit(sid)


			return JsonResponse({'message':'El registro se ha guardado correctamente','success':'ok',
					'data':''})
			
		except Exception,e:
			#print e
			transaction.savepoint_rollback(sid)
			return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
				'data':''})




#eliminar la asociacion de circuito lote, pagina editar
@transaction.atomic
def eliminarAsociacionCircuitoLote(request):

	sid = transaction.savepoint()
	try:
		lista=request.POST['_content']
		respuesta= json.loads(lista)
		
		for item in respuesta['lista']:
			LoteCircuito.objects.filter(id=item['id']).delete()

		transaction.savepoint_commit(sid)
		return JsonResponse({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''})

	except ProtectedError:
		return JsonResponse({'message':'No es posible eliminar el registro, se esta utilizando en otra seccion del sistema','success':'error','data':''})
		
	except Exception,e:
		#print e
		transaction.savepoint_rollback(sid)
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})		





#consultar sector, pagina editar
@api_view(['GET'])
def consulta_listado_sector(request):

	if request.method == 'GET':

		try:
			qset=''
			dato= request.GET['dato']
			id_lote= request.GET['id_lote']

			qset = Q(activo=True)

			if dato:
					qset =  qset &(Q(nombre__icontains=dato))


			sector = Sector.objects.filter(qset)

			lista_sector=[]

			for item in list(sector):

			 	sector_lote = LoteSector.objects.filter(sector__id=item.id,lote__id=id_lote).first()

			 	if sector_lote is None:

					lista_sector.append(

						{		
							'id':item.id,
							'nombre':item.nombre,
						}

					)

			return JsonResponse({'message':'','success':'ok','data':lista_sector})

		except Exception,e:
			print e
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

	    #return response



#consultar sector asociados al lote, pagina editar
@api_view(['GET'])
def consulta_listado_sector_asociados(request):

	if request.method == 'GET':

		try:

			id_lote= request.GET['id_lote']

			sector__asociado_lote = LoteSector.objects.filter(lote__id=id_lote)

			lista_sector_asociado=[]

			for item in list(sector__asociado_lote):

				lista_sector_asociado.append(

					{		
						'id':item.id,
						'nombre':item.sector.nombre,
					}

				)

			return JsonResponse({'message':'','success':'ok','data':lista_sector_asociado})

		except Exception,e:
			print e
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

    #return response



#guardar la asociacion de los sectores al lote pagina editar
@api_view(['POST'])
@transaction.atomic
def AsociarSectorLoteEditar(request):

	sid = transaction.savepoint()
	if request.method == 'POST':
		try:
			
			lista=request.DATA['lista']
			lote_id= request.DATA['lote_id']
  
			for item in lista:

				circuitoLote=LoteSector(sector_id=item['id'],lote_id=lote_id)
				circuitoLote.save()

				transaction.savepoint_commit(sid)


			return JsonResponse({'message':'El registro se ha guardado correctamente','success':'ok',
					'data':''})
			
		except Exception,e:
			#print e
			transaction.savepoint_rollback(sid)
			return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
				'data':''})




#eliminar la asociacion de sector lote, pagina editar
@transaction.atomic
def eliminarAsociacionSectorLote(request):

	sid = transaction.savepoint()
	try:
		lista=request.POST['_content']
		respuesta= json.loads(lista)
		
		for item in respuesta['lista']:
			LoteSector.objects.filter(id=item['id']).delete()

		transaction.savepoint_commit(sid)
		return JsonResponse({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''})

	except ProtectedError:
		return JsonResponse({'message':'No es posible eliminar el registro, se esta utilizando en otra seccion del sistema','success':'error','data':''})
		
	except Exception,e:
		#print e
		transaction.savepoint_rollback(sid)
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})





#guardar la asociacion de los poligonos al lote pagina editar
@api_view(['POST'])
@transaction.atomic
def AsociarPoligonoLoteEditar(request):

	sid = transaction.savepoint()
	if request.method == 'POST':
		try:
			
			lista=request.DATA['lista']
			lote_id= request.DATA['lote_id']
  
			for item in lista:

				print item['nombre']

				poligono=Poligono(nombre=item['nombre'],activo=True)
				poligono.save()


				poligono_lote=LotePoligono(poligono_id=poligono.id,lote_id=lote_id)
				poligono_lote.save()

				transaction.savepoint_commit(sid)


			return JsonResponse({'message':'El registro se ha guardado correctamente','success':'ok',
					'data':''})
			
		except Exception,e:
			#print e
			transaction.savepoint_rollback(sid)
			return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
				'data':''})




#consultar sector asociados al poliogno, pagina editar
@api_view(['GET'])
def consulta_listado_poligono_asociados(request):

	if request.method == 'GET':

		try:

			id_lote= request.GET['id_lote']

			poligono__asociado_lote = LotePoligono.objects.filter(lote__id=id_lote)

			lista_poligono_asociado=[]

			for item in list(poligono__asociado_lote):

				lista_poligono_asociado.append(

					{		
						'id':item.id,
						'nombre':item.poligono.nombre,
					}

				)

			return JsonResponse({'message':'','success':'ok','data':lista_poligono_asociado})

		except Exception,e:
			print e
			return JsonResponse({'message':'Se presentaron errores de comunicacion con el servidor','status':'error','data':''},status=status.HTTP_500_INTERNAL_SERVER_ERROR)		

    #return response



#eliminar la asociacion de sector lote, pagina editar
@transaction.atomic
def eliminarAsociacionPoligonoLote(request):

	sid = transaction.savepoint()
	try:
		lista=request.POST['_content']
		respuesta= json.loads(lista)
		
		for item in respuesta['lista']:

			lote_poligono = LotePoligono.objects.filter(id=item['id']).first()

			LotePoligono.objects.filter(id=item['id']).delete()
			Poligono.objects.filter(id=lote_poligono.poligono.id).delete()

		transaction.savepoint_commit(sid)
		return JsonResponse({'message':'El registro se ha eliminado correctamente','success':'ok',
				'data':''})

	except ProtectedError:
		return JsonResponse({'message':'No es posible eliminar el registro, se esta utilizando en otra seccion del sistema','success':'error','data':''})
		
	except Exception,e:
		#print e
		transaction.savepoint_rollback(sid)
		return JsonResponse({'message':'Se presentaron errores al procesar la solicitud','success':'error',
			'data':''})





def EjemploWord(request):

	ano= request.GET['ano']
	mes= request.GET['mes']
	inspeccion= request.GET['inspeccion']
	lote_id= request.GET['lote_id']
	qset=''


	qset = Q(lote__id=lote_id)

	if inspeccion:
		qset =  qset &(Q(numero_inspeccion=inspeccion))


	if ano:
		qset =  qset &(Q(fecha__year=ano))


	if mes:
		qset =  qset &(Q(fecha__month=mes))

	
	#inspeccion = AInspeccion.objects.filter(qset).values('sector__id','sector__nombre').distinct()
	#listsector=''
	# for item in list(inspeccion):

	# 	listsector='{}{},'.format(listsector,item['sector__nombre'])

	# print listsector

	inspeccion = AInspeccion.objects.filter(qset).first()

	newpath = r'media/prueba/'
	functions.descargarArchivoS3(str('plantillas/SCP/informeSCP.docx'),str(newpath))

	document = Document(newpath+'informeSCP.docx')

	# styles = document.styles

	# for item in styles:
	# 	print item

	# font = styles['Normal'].font
	# font.name = 'Arial'

	# sections = document.sections
	# section = sections[0]
	# print section.hearder
	
	#document.add_heading('Document Title', 0)
	

	encabezado = document.add_paragraph('')
	parrafo = encabezado.add_run(unicode('INFORME DE INSPECCIÓN DE CALIDAD DE LAS OBRAS', 'utf-8'))
	parrafo.bold = True
	parrafo.italic=False
	parrafo.font.size = Pt(22)
	
	p = document.add_paragraph('')
	parrafo2=p.add_run('___________________________________________')
	parrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	parrafo2.font.size = Pt(20)


	encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER

	#linea = document.add_paragraph('',style='IntenseQuote')
	linea = document.add_paragraph('')

	fecha=str(inspeccion.fecha).split('-')



	p = document.add_paragraph('PROYECTO: '+ inspeccion.apoyo.nombre)
	p = document.add_paragraph('CIRCUITO: '+ inspeccion.circuito.nombre)
	p = document.add_paragraph('DISTRIBUIDORA: '+ inspeccion.lote.sucursal.nombre)
	p = document.add_paragraph('CONTRATISTA: '+ inspeccion.lote.contratista.nombre)
	p = document.add_paragraph(unicode('INSPECCIÓN: ', 'utf-8')+ str(inspeccion.numero_inspeccion) +' DEL MES DE '+ str(fecha[1]) +' DE '+ str(fecha[0]))
	p = document.add_paragraph('SECTOR: '+ inspeccion.sector.nombre)
	p = document.add_paragraph('INSPECTORES FIRMA VERIFICADORA CUSA ATRATO: ')

	# p.add_run('bold').bold = True
	# p.add_run(' and some ')
	# p.add_run('italic.').italic = True

	p = document.add_paragraph('') #parrafo vacio (espacio enter)

	#observacion
	#p1 = document.add_paragraph('',style='ListNumber')
	p = document.add_paragraph('')
	observ = p.add_run('1. OBJETIVO.')
	observ.bold = True
	p = document.add_paragraph(unicode('Realizar un acompañamiento durante el avance del proyecto para retroalimentar sobre posibles errores en la construcción entregando las recomendaciones particulares sobre las mejoras que permitan prever o corregir las desviaciones en las actividades ejecutadas.', 'utf-8'))


	p = document.add_paragraph('') #parrafo vacio (espacio enter)

	#descripcion
	#p3 = document.add_paragraph('',style='ListNumber')
	p = document.add_paragraph('')
	desc = p.add_run(unicode('2. DESCRIPCIÓN.', 'utf-8'))
	desc.bold = True

	#p = document.add_paragraph(unicode('Se realizó inspección al sector Barrio Nuevo donde desarrollan las obras del proyecto lote 5 OFID, circuito NIBA-112 polígono 2. En este sector se ha realizado el izaje de postes, con cimentaciones, instalación de puestas a tierra, instalación de armados, estructuras conectadas al neutro del sistema y de puestas a tierra a través del cable de tierra del poste, tendidos de cable de media y baja tensión e instalación de transformadores.', 'utf-8'))



	p = document.add_paragraph('') #parrafo vacio (espacio enter)

	#DESARROLLO DE LA INSPECCIoN
	#p6 = document.add_paragraph('',style='ListNumber')
	p = document.add_paragraph('')
	descInps = p.add_run(unicode('3. DESARROLLO DE LA INSPECCIÓN.', 'utf-8'))
	descInps.bold = True

	#p = document.add_paragraph(unicode('En los apoyos inspeccionados se verificaron para el izaje de postes la verticalidad o aplomo de los postes instalados, la instalación de las tierras, cimentaciones y armados de estructuras, puestas a tierra de las estructuras, la instalación de los vientos, la correcta instalación de transformadores, cierre del neutro del sistema y del cable de tierra en los postes con su varilla de cobre, tendidos de cables de media y baja tensión.', 'utf-8'))

	p = document.add_paragraph('') #parrafo vacio (espacio enter)

	# document.add_heading('Heading, level 1', level=1)
	# document.add_paragraph('Intense quote', style='IntenseQuote')

	# document.add_paragraph(
	#     'first item in unordered list', style='ListBullet'
	# )

	#document.add_picture('monty-truth.png', width=Inches(1.25))


	falla_inspeccion = BFallaInspeccion.objects.filter(inspeccion__id=inspeccion.id)

	#print inspeccion.id

	for item in list(falla_inspeccion):

		foto_falla = FotoFallaInspeccion.objects.filter(falla_inspeccion__id=item.id)
		
		valor=0
		calificaciones=''
		for item2 in list(foto_falla):

			# print item2.id
			# print item2.soporte
			#a=document.add_picture(item2.soporte, width=Inches(1.25))

			if int(item2.falla_inspeccion.calificacion)==1:

				calificaciones='Conforme'

			elif int(item2.falla_inspeccion.calificacion)==2:

				calificaciones='No Conforme'

			elif int(item2.falla_inspeccion.calificacion)==3:

				calificaciones='No aplica'


			if valor==0:				
				table = document.add_table(rows=2, cols=2)
				table.style ='Table Grid'
			
			
			hdr_cells = table.rows[0].cells

			paragraph = hdr_cells[valor].paragraphs[0]
			run = paragraph.add_run()
			run.add_picture(item2.soporte, width = 2800000, height = 1600000)

			
			row_cells = table.rows[1].cells
			#row_cells[0].text = str('dsadasda')
			row_cells[valor].text = str('Postes '+ item2.falla_inspeccion.inspeccion.apoyo.nombre + ' ' +item2.falla_inspeccion.capitulo_falla.descripcion+'. '+calificaciones)

			valor=valor+1

			if valor==2:
				p= document.add_paragraph('') #parrafo vacio (espacio enter)
				addTableAfterParagraph(table,p)
				valor=0

			#row_cells[2].text = '2'
	

	p = document.add_paragraph('') #parrafo vacio (espacio enter)

	#RECOMENDACIONES
	#p6 = document.add_paragraph('',style='ListNumber')
	p = document.add_paragraph('')
	descInps = p.add_run('4. RECOMENDACIONES.')
	descInps.bold = True
		

	document.add_page_break()

	nombreArchivo='INF DE CALIDAD.docx'

	document.save(nombreArchivo)


	chunk_size = 108192

	response = StreamingHttpResponse(FileWrapper(open(nombreArchivo,'rb'),chunk_size),
				content_type=mimetypes.guess_type(nombreArchivo)[0])
	response['Content-Length'] = os.path.getsize(nombreArchivo)
	response['Content-Disposition'] = "attachment; filename=%s" % nombreArchivo
			
	return response



def addTableAfterParagraph(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


@login_required
def Lotes(request):
	qsContratista=Contratista.objects.all()
	qsSector=Sector.objects.filter(activo=True)
	qsCircuito=Circuito.objects.filter(activo=True)
	qsmes=Mes.objects.all()


	return render_to_response('lotes/lote.html',{'mes':qsmes,'contratista':qsContratista,'sector':qsSector,'circuito':qsCircuito,'app':'lote','model':'Lote'},context_instance=RequestContext(request))

#registrar el lote
@login_required
def RegistroLote(request):
	
	qscontratista=Contratista.objects.all()
	qsprovincia=Provincia.objects.all()
	qssucursal=Sucursal.objects.filter(activo=True)

	return render_to_response('lotes/registro_lote.html',{'sucursal':qssucursal,'provincia':qsprovincia,'contratista':qscontratista,'app':'lote','model':'Lote'},context_instance=RequestContext(request))



@login_required
def ActualizarLote(request,id_lote=None):
	
	qscontratista=Contratista.objects.all()
	qsprovincia=Provincia.objects.all()
	qssucursal=Sucursal.objects.filter(activo=True)

	return render_to_response('lotes/editar_lote.html',{'sucursal':qssucursal,'provincia':qsprovincia,'contratista':qscontratista,'app':'lote','model':'Lote','id_lote':id_lote},context_instance=RequestContext(request))
